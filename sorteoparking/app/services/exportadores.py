"""Exportación de actas del sorteo. SDD §16, T-110, T-119."""

import io
import hashlib
from datetime import datetime, timezone
from typing import Any

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from sqlalchemy.orm import Session

from app.models.catalogo import Zona
from app.models.sorteo import (
    Consejero,
    Participante,
    ResultadoSorteo,
    SesionOTP,
    Sorteo,
)
from app.models.tenant import Tenant
from app.models.log import LogAuditoria


# SDD §16.7 — Protección contra Formula Injection
FORMULA_PREFIXES = ("=", "+", "-", "@", "\t", "\r")


def _sanitizar(valor: Any) -> str:
    """Previene Formula Injection en celdas Excel."""
    if not isinstance(valor, str):
        valor = str(valor) if valor is not None else ""
    if valor.startswith(FORMULA_PREFIXES):
        return "'" + valor
    return valor


def _estilo_encabezado():
    return {
        "font": Font(bold=True, color="FFFFFF", size=11),
        "fill": PatternFill(start_color="1A3A5C", end_color="1A3A5C", fill_type="solid"),
        "alignment": Alignment(horizontal="center", vertical="center", wrap_text=True),
        "border": Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        ),
    }


def _estilo_celda():
    return {
        "alignment": Alignment(vertical="center", wrap_text=False),
        "border": Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        ),
    }


def exportar_acta_excel(db: Session, tenant_id: str, sorteo_id: int) -> bytes:
    """Genera acta Excel con 5 hojas (SDD §16.5)."""
    sorteo = db.query(Sorteo).filter(
        Sorteo.id == sorteo_id, Sorteo.tenant_id == tenant_id
    ).first()
    if not sorteo:
        raise ValueError("Sorteo no encontrado")

    tenant = db.query(Tenant).filter(Tenant.id == tenant_id).first()

    wb = Workbook()

    # ---- Hoja 1: Resumen ----
    ws1 = wb.active
    ws1.title = "Resumen"
    ws1.merge_cells("A1:B1")
    ws1["A1"] = f"ACTA DE SORTEO — {_sanitizar(tenant.nombre if tenant else '')}"
    ws1["A1"].font = Font(bold=True, size=14, color="1A3A5C")

    data_resumen = [
        ("Conjunto", tenant.nombre if tenant else ""),
        ("Fecha", sorteo.created_at.strftime("%Y-%m-%d %H:%M UTC") if sorteo.created_at else ""),
        ("Tipo", sorteo.tipo or "GENERAL"),
        ("Estado", sorteo.estado),
        ("Modelo aplicado", sorteo.modelo_aplicado or "No especificado"),
        ("Seed (SHA-256)", sorteo.seed or ""),
        ("Snapshot hash", sorteo.snapshot_hash or ""),
    ]
    for i, (k, v) in enumerate(data_resumen, start=3):
        ws1[f"A{i}"] = _sanitizar(k)
        ws1[f"A{i}"].font = Font(bold=True)
        ws1[f"B{i}"] = _sanitizar(v)

    # Totales
    total_part = db.query(Participante).filter(
        Participante.sorteo_id == sorteo_id, Participante.tenant_id == tenant_id
    ).count()
    ganadores = db.query(ResultadoSorteo).filter(
        ResultadoSorteo.sorteo_id == sorteo_id,
        ResultadoSorteo.tenant_id == tenant_id,
        ResultadoSorteo.tipo_resultado == "GANADOR",
    ).count()
    perdedores = db.query(ResultadoSorteo).filter(
        ResultadoSorteo.sorteo_id == sorteo_id,
        ResultadoSorteo.tenant_id == tenant_id,
        ResultadoSorteo.tipo_resultado == "PERDEDOR",
    ).count()

    row = len(data_resumen) + 4
    ws1[f"A{row}"] = "Total participantes"
    ws1[f"A{row}"].font = Font(bold=True)
    ws1[f"B{row}"] = total_part
    ws1[f"A{row+1}"] = "Ganadores"
    ws1[f"A{row+1}"].font = Font(bold=True)
    ws1[f"B{row+1}"] = ganadores
    ws1[f"A{row+2}"] = "No asignados"
    ws1[f"A{row+2}"].font = Font(bold=True)
    ws1[f"B{row+2}"] = perdedores

    ws1.column_dimensions["A"].width = 25
    ws1.column_dimensions["B"].width = 70

    # ---- Hoja 2: Ganadores ----
    ws2 = wb.create_sheet("Ganadores")
    headers = ["Apartamento", "Parqueadero", "Zona", "Tipo Vehículo", "Reasignado"]
    for col, h in enumerate(headers, start=1):
        cell = ws2.cell(row=1, column=col, value=_sanitizar(h))
        for k, v in _estilo_encabezado().items():
            setattr(cell, k, v)

    resultados = (
        db.query(ResultadoSorteo, Participante)
        .join(Participante, ResultadoSorteo.participante_id == Participante.id)
        .filter(
            ResultadoSorteo.sorteo_id == sorteo_id,
            ResultadoSorteo.tenant_id == tenant_id,
            ResultadoSorteo.tipo_resultado == "GANADOR",
        )
        .order_by(Participante.apartamento)
        .all()
    )
    for i, (r, p) in enumerate(resultados, start=2):
        ws2.cell(row=i, column=1, value=_sanitizar(p.apartamento))
        ws2.cell(row=i, column=2, value=_sanitizar(r.parqueadero_asignado))
        ws2.cell(row=i, column=3, value=_sanitizar(r.zona_asignada))
        ws2.cell(row=i, column=4, value=_sanitizar(p.tipo_vehiculo))
        ws2.cell(row=i, column=5, value="Sí" if r.fue_reasignado else "No")

    for col in range(1, len(headers) + 1):
        ws2.column_dimensions[get_column_letter(col)].width = 20

    # ---- Hoja 3: No Asignados ----
    ws3 = wb.create_sheet("No Asignados")
    headers3 = ["Apartamento", "Tipo Vehículo"]
    for col, h in enumerate(headers3, start=1):
        cell = ws3.cell(row=1, column=col, value=_sanitizar(h))
        for k, v in _estilo_encabezado().items():
            setattr(cell, k, v)

    no_asignados = (
        db.query(ResultadoSorteo, Participante)
        .join(Participante, ResultadoSorteo.participante_id == Participante.id)
        .filter(
            ResultadoSorteo.sorteo_id == sorteo_id,
            ResultadoSorteo.tenant_id == tenant_id,
            ResultadoSorteo.tipo_resultado == "PERDEDOR",
        )
        .order_by(Participante.apartamento)
        .all()
    )
    for i, (r, p) in enumerate(no_asignados, start=2):
        ws3.cell(row=i, column=1, value=_sanitizar(p.apartamento))
        ws3.cell(row=i, column=2, value=_sanitizar(p.tipo_vehiculo))

    for col in range(1, len(headers3) + 1):
        ws3.column_dimensions[get_column_letter(col)].width = 20

    # ---- Hoja 4: Consejeros ----
    ws4 = wb.create_sheet("Consejeros")
    headers4 = ["Nombre", "Email", "Estado OTP", "Confirmado en"]
    for col, h in enumerate(headers4, start=1):
        cell = ws4.cell(row=1, column=col, value=_sanitizar(h))
        for k, v in _estilo_encabezado().items():
            setattr(cell, k, v)

    consejeros = (
        db.query(SesionOTP, Consejero)
        .join(Consejero, SesionOTP.consejero_id == Consejero.id)
        .filter(SesionOTP.sorteo_id == sorteo_id, SesionOTP.tenant_id == tenant_id)
        .order_by(Consejero.id)
        .all()
    )
    for i, (ses, cons) in enumerate(consejeros, start=2):
        ws4.cell(row=i, column=1, value=_sanitizar(cons.nombre))
        # Los consejeros aparecen con sus datos según SDD §16.2
        ws4.cell(row=i, column=2, value=_sanitizar(cons.email))
        ws4.cell(row=i, column=3, value=_sanitizar(ses.estado))
        ts = ses.confirmado_en.strftime("%Y-%m-%d %H:%M:%S UTC") if ses.confirmado_en else ""
        ws4.cell(row=i, column=4, value=ts)

    for col in range(1, len(headers4) + 1):
        ws4.column_dimensions[get_column_letter(col)].width = 25

    # ---- Hoja 5: Log Auditoría ----
    ws5 = wb.create_sheet("Log Auditoria")
    headers5 = ["Evento", "Payload", "Hash Anterior", "Hash Actual", "Timestamp"]
    for col, h in enumerate(headers5, start=1):
        cell = ws5.cell(row=1, column=col, value=_sanitizar(h))
        for k, v in _estilo_encabezado().items():
            setattr(cell, k, v)

    logs = (
        db.query(LogAuditoria)
        .filter(LogAuditoria.tenant_id == tenant_id)
        .order_by(LogAuditoria.id)
        .limit(500)  # limit to prevent memory issues
        .all()
    )
    for i, log in enumerate(logs, start=2):
        ws5.cell(row=i, column=1, value=_sanitizar(log.evento))
        ws5.cell(row=i, column=2, value=_sanitizar(log.payload))
        ws5.cell(row=i, column=3, value=_sanitizar(log.hash_anterior))
        ws5.cell(row=i, column=4, value=_sanitizar(log.hash_actual))
        ws5.cell(row=i, column=5, value=log.created_at.strftime("%Y-%m-%d %H:%M:%S UTC") if log.created_at else "")

    ws5.column_dimensions["A"].width = 30
    ws5.column_dimensions["B"].width = 45
    ws5.column_dimensions["C"].width = 70
    ws5.column_dimensions["D"].width = 70
    ws5.column_dimensions["E"].width = 25

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def exportar_acta_word(db: Session, tenant_id: str, sorteo_id: int) -> bytes:
    """Genera acta Word (SDD §16.4)."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    sorteo = db.query(Sorteo).filter(
        Sorteo.id == sorteo_id, Sorteo.tenant_id == tenant_id
    ).first()
    if not sorteo:
        raise ValueError("Sorteo no encontrado")
    tenant = db.query(Tenant).filter(Tenant.id == tenant_id).first()

    doc = Document()

    # Título
    title = doc.add_heading(f"ACTA DE SORTEO — {_sanitizar(tenant.nombre if tenant else '')}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Datos generales
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Datos del sorteo")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    datos = [
        ("Conjunto", tenant.nombre if tenant else ""),
        ("Fecha", sorteo.created_at.strftime("%Y-%m-%d %H:%M UTC") if sorteo.created_at else ""),
        ("Tipo", sorteo.tipo or "GENERAL"),
        ("Seed", sorteo.seed or ""),
    ]
    for k, v in datos:
        p = doc.add_paragraph()
        run_k = p.add_run(f"{k}: ")
        run_k.bold = True
        p.add_run(_sanitizar(v))

    # Consejeros
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Consejeros garantes")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    consejeros = (
        db.query(SesionOTP, Consejero)
        .join(Consejero, SesionOTP.consejero_id == Consejero.id)
        .filter(SesionOTP.sorteo_id == sorteo_id, SesionOTP.tenant_id == tenant_id)
        .order_by(Consejero.id)
        .all()
    )
    for ses, cons in consejeros:
        ts = ses.confirmado_en.strftime("%Y-%m-%d %H:%M:%S UTC") if ses.confirmado_en else "Pendiente"
        doc.add_paragraph(f"• {_sanitizar(cons.nombre)} — OTP: {ses.estado} — {ts}", style="List Bullet")

    # Resultados
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Resultados")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    ganadores = (
        db.query(ResultadoSorteo, Participante)
        .join(Participante, ResultadoSorteo.participante_id == Participante.id)
        .filter(
            ResultadoSorteo.sorteo_id == sorteo_id,
            ResultadoSorteo.tenant_id == tenant_id,
            ResultadoSorteo.tipo_resultado == "GANADOR",
        )
        .order_by(Participante.apartamento)
        .all()
    )
    no_asignados_count = db.query(ResultadoSorteo).filter(
        ResultadoSorteo.sorteo_id == sorteo_id,
        ResultadoSorteo.tenant_id == tenant_id,
        ResultadoSorteo.tipo_resultado == "PERDEDOR",
    ).count()

    doc.add_paragraph(f"Total ganadores: {len(ganadores)}")
    doc.add_paragraph(f"No asignados: {no_asignados_count}")

    # Seed destacado
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Seed de verificación")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    p = doc.add_paragraph()
    run = p.add_run(_sanitizar(sorteo.seed or "No disponible"))
    run.font.size = Pt(10)
    run.font.name = "Courier New"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
